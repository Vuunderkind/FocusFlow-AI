"""
Генератор курсовой работы для FocusFlow AI
Запуск: py generate_coursework.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Поля страницы ────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(1.5)

# ─── Цвета ────────────────────────────────────────────────────────
PURPLE  = RGBColor(0x7C, 0x5C, 0xFC)
DARK    = RGBColor(0x1C, 0x1B, 0x22)
GRAY    = RGBColor(0x44, 0x44, 0x66)
BLACK   = RGBColor(0x00, 0x00, 0x00)
HEADING_COLOR = RGBColor(0x5C, 0x38, 0xCC)

# ─── Вспомогательные функции ──────────────────────────────────────
def add_heading(text, level=1, color=HEADING_COLOR):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_term(term, definition):
    """Добавить термин с определением"""
    p = doc.add_paragraph()
    r1 = p.add_run(f"📌 {term} ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = PURPLE
    r1.font.name = 'Segoe UI'
    r2 = p.add_run(f"— {definition}")
    r2.font.size = Pt(11)
    r2.font.name = 'Segoe UI'
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x10, 0x60, 0x10)
    # Gray background hack via shading
    try:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
    except: pass
    return p

def add_separator():
    p = doc.add_paragraph('─' * 65)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xDD)
    run.font.size = Pt(8)

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  ТИТУЛЬНЫЙ ЛИСТ
# ══════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ХЕКСЛЕТ КОЛЛЕДЖ")
r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'

add_body("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Направление: Программная инженерия (ПМ3)")
r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("КУРСОВАЯ РАБОТА")
r.bold = True; r.font.size = Pt(22); r.font.name = 'Times New Roman'
r.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("по теме:")
r.font.size = Pt(13); r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("«Автоматизация и объединение ИИ»")
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0x33, 0x22, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Проект: FocusFlow AI")
r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Windows-приложение для интеллектуального общения с ИИ\n"
    "и автоматизации задач через оркестрацию языковых моделей"
)
r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.italic = True

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Выполнил: студент группы ПМ3\nГод: 2025–2026")
r.font.size = Pt(12); r.font.name = 'Times New Roman'

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  СОДЕРЖАНИЕ
# ══════════════════════════════════════════════════════════════════

add_heading("СОДЕРЖАНИЕ", 1)
contents = [
    ("1.", "Введение. Что такое FocusFlow AI?"),
    ("2.", "СЛОВАРЬ ТЕРМИНОВ"),
    ("3.", "БЭКЕНД (Backend) — серверная и логическая часть"),
    ("  3.1", "Язык и платформа"),
    ("  3.2", "Модели данных (Models)"),
    ("  3.3", "База данных и репозитории (Data)"),
    ("  3.4", "Сервисы и бизнес-логика (Services)"),
    ("  3.5", "Провайдеры ИИ"),
    ("  3.6", "Маршрутизатор запросов (AIRouter)"),
    ("  3.7", "Оркестратор (OrchestrationService)"),
    ("  3.8", "Воркфлоу — цепочки агентов (WorkflowService)"),
    ("4.", "ФРОНТЕНД (Frontend) — визуальный интерфейс"),
    ("  4.1", "Дизайн-система (Theme.cs)"),
    ("  4.2", "Переиспользуемые компоненты (Controls)"),
    ("  4.3", "Главное окно (MainForm)"),
    ("  4.4", "Страница чата (ChatPage)"),
    ("  4.5", "Страница агентов (AgentsPage)"),
    ("  4.6", "Страница автоматизации (AutomationPage)"),
    ("  4.7", "История чатов (HistoryPage)"),
    ("  4.8", "Настройки (SettingsPage)"),
    ("5.", "Как всё работает вместе: полный цикл запроса"),
    ("6.", "Структура файлов проекта"),
    ("7.", "Заключение"),
]
for num, title in contents:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  1. ВВЕДЕНИЕ
# ══════════════════════════════════════════════════════════════════

add_heading("1. ВВЕДЕНИЕ. ЧТО ТАКОЕ FOCUSFLOW AI?", 1)

add_body(
    "FocusFlow AI — это настольное Windows-приложение, написанное на языке C# (.NET 10). "
    "Программа позволяет пользователю общаться с несколькими искусственными интеллектами "
    "(Claude, GPT, Gemini, Mistral, Groq) через единый красивый интерфейс. Помимо обычного "
    "чата, приложение умеет автоматически выбирать лучший ИИ для каждого вопроса, объединять "
    "ответы двух моделей в один идеальный ответ, а также запускать цепочки автоматических "
    "задач (воркфлоу).", indent=True
)

add_body(
    "Проект был создан в рамках учебной практики по теме «Автоматизация и объединение ИИ» "
    "в Хекслет Колледже.", indent=True
)

add_heading("Основные возможности:", 2)
features = [
    "💬  Чат с ИИ — задавайте вопросы любым AI-моделям",
    "✦   AUTO-режим — программа сама выбирает лучший ИИ для вашего вопроса",
    "⚡  FUSION-режим — два разных ИИ отвечают параллельно, затем синтезатор объединяет лучшее",
    "🤖  Агенты — специализированные помощники (Программист, Аналитик, Дизайнер и др.)",
    "🔁  Воркфлоу — цепочки задач, где выход одного агента становится входом следующего",
    "💾  История — все разговоры сохраняются в локальную базу данных SQLite",
    "⚙️  Настройки — хранение API-ключей всех провайдеров в зашифрованном конфиге",
]
for f in features:
    p = doc.add_paragraph(f, style='List Bullet')
    p.runs[0].font.size = Pt(11)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  2. СЛОВАРЬ ТЕРМИНОВ
# ══════════════════════════════════════════════════════════════════

add_heading("2. СЛОВАРЬ ТЕРМИНОВ", 1)
add_body("Ниже приведены все технические термины, встречающиеся в этой работе.")

terms = [
    ("API (Application Programming Interface)",
     "интерфейс программирования — набор правил, по которым одна программа общается с другой. "
     "В нашем случае: правила, по которым мы отправляем запрос на сервер ИИ и получаем ответ."),
    ("API-ключ (API Key)",
     "уникальный пароль/токен, который выдаёт сервис (OpenAI, Anthropic и т.д.), "
     "чтобы идентифицировать вашу программу и списывать оплату с вашего аккаунта."),
    ("Бэкенд (Backend)",
     "«невидимая» серверная/логическая часть программы. Здесь обрабатываются данные, "
     "выполняются вычисления, происходит сохранение в базу данных. Пользователь это не видит, "
     "но без этого ничего не работает."),
    ("Фронтенд (Frontend)",
     "видимая часть программы — кнопки, поля ввода, окна, цвета, шрифты. "
     "Всё, с чем пользователь непосредственно взаимодействует."),
    ("C# (си шарп)",
     "язык программирования от Microsoft. Используется для создания Windows-приложений, "
     "серверных систем, игр (Unity). Наш проект написан полностью на C#."),
    (".NET (dot net)",
     "платформа/фреймворк от Microsoft, на которой работают программы, написанные на C#. "
     "Версия 10 — самая современная на момент разработки проекта."),
    ("Windows Forms (WinForms)",
     "библиотека для создания оконных приложений в Windows. Именно она рисует окна, "
     "кнопки, поля ввода. Это наш «фронтенд-движок»."),
    ("SQLite",
     "лёгкая встроенная база данных — один файл на диске, не требует отдельного сервера. "
     "Идеально для настольных приложений. Мы храним в ней чаты, агентов, воркфлоу."),
    ("База данных (Database, БД)",
     "организованное хранилище информации. Как большая умная таблица Excel, "
     "только быстрее и надёжнее."),
    ("Таблица (Table)",
     "структура в базе данных, аналог листа Excel. У нас есть таблицы: "
     "Conversations (чаты), Messages (сообщения), Agents (агенты), Workflows (воркфлоу)."),
    ("SQL (Structured Query Language)",
     "язык запросов к базе данных. Примеры: SELECT — выбрать данные, "
     "INSERT — добавить, UPDATE — обновить, DELETE — удалить."),
    ("Репозиторий (Repository)",
     "класс, который берёт на себя всю работу с базой данных. "
     "Вместо того чтобы писать SQL-запросы везде, мы вызываем методы репозитория: "
     "например, ConversationRepository.GetAll() — получить все чаты."),
    ("Класс (Class)",
     "шаблон/чертёж объекта в программировании. Например, класс AIAgent описывает, "
     "какие поля есть у агента (имя, описание, системный промпт и т.д.)."),
    ("Объект (Object)",
     "конкретный экземпляр класса. Если класс — это чертёж стула, то объект — это конкретный стул."),
    ("Интерфейс (Interface)",
     "в программировании — контракт, который обязует класс реализовать определённые методы. "
     "Например, ILLMProvider говорит: «любой провайдер ИИ ОБЯЗАН иметь метод SendAsync»."),
    ("Метод (Method)",
     "функция внутри класса. Например, метод SendAsync у каждого провайдера отправляет "
     "запрос к ИИ и возвращает ответ."),
    ("Async/Await (асинхронность)",
     "механизм, позволяющий программе не «замирать» пока ИИ думает. "
     "Пока ждём ответа от сервера, интерфейс остаётся отзывчивым."),
    ("HTTP-запрос",
     "способ общения между программами через интернет. "
     "Как отправить письмо (запрос) и получить ответ. Мы отправляем POST-запросы к API ИИ."),
    ("JSON (JavaScript Object Notation)",
     "формат хранения и передачи данных. Выглядит как: {\"role\": \"user\", \"content\": \"Привет\"}. "
     "В этом формате мы отправляем запросы к ИИ и получаем ответы."),
    ("LLM (Large Language Model)",
     "большая языковая модель — вид ИИ, обученный на огромных текстовых массивах. "
     "Умеет отвечать на вопросы, писать код, переводить тексты. Примеры: Claude, GPT-4, Gemini."),
    ("Провайдер (Provider)",
     "компания, которая предоставляет доступ к ИИ через API. "
     "Anthropic (Claude), OpenAI (GPT), Google (Gemini), Mistral, Groq."),
    ("Промпт (Prompt)",
     "сообщение/инструкция, отправляемая ИИ. Системный промпт — базовые инструкции "
     "для агента (например: «Ты — эксперт по программированию»)."),
    ("Токен (Token)",
     "единица текста для ИИ. Примерно 1 токен = 3/4 слова. "
     "Оплата API идёт за токены: чем длиннее диалог — тем дороже."),
    ("Оркестрация (Orchestration)",
     "управление несколькими ИИ-моделями — выбор правильной, координация параллельной "
     "работы, объединение результатов."),
    ("Маршрутизатор (Router)",
     "компонент, который анализирует запрос и решает: какой ИИ лучше справится. "
     "Как GPS — выбирает оптимальный маршрут."),
    ("FUSION-режим",
     "режим, при котором два разных ИИ отвечают на один вопрос одновременно, "
     "а третий ИИ (синтезатор) объединяет лучшее из обоих ответов."),
    ("Воркфлоу (Workflow)",
     "цепочка задач, где результат одного шага становится входными данными следующего. "
     "Например: Шаг1 (ИИ-исследователь пишет текст) → Шаг2 (ИИ-редактор правит) → "
     "Шаг3 (ИИ-переводчик переводит)."),
    ("Агент (AI Agent)",
     "специализированный ИИ-помощник с конкретной ролью и системным промптом. "
     "В нашем проекте: Code Master, Content Writer, Data Analyst и другие."),
    ("Событие (Event)",
     "механизм в Windows Forms: когда пользователь нажимает кнопку — "
     "возникает событие Click, которое запускает нужный код."),
    ("UserControl",
     "переиспользуемый визуальный компонент WinForms. Как «мини-форма» внутри главного окна. "
     "Каждая страница приложения (Чат, Агенты, История) — это UserControl."),
    ("DPI (Dots Per Inch)",
     "плотность пикселей экрана. High-DPI = экраны с высоким разрешением (Retina, 4K). "
     "Программа учитывает это, чтобы интерфейс выглядел одинаково хорошо."),
    ("Конфиг (Config)",
     "файл настроек программы. В нашем случае: config.json хранит API-ключи "
     "и пользовательские предпочтения."),
    ("AppData",
     "системная папка Windows (%AppData%), где программы хранят свои данные. "
     "Наша БД и конфиг лежат в C:\\Users\\Имя\\AppData\\Roaming\\FocusFlowAI\\"),
    ("Сериализация (Serialization)",
     "преобразование объекта C# в строку JSON для сохранения или отправки. "
     "Десериализация — обратный процесс: JSON → объект C#."),
    ("Gradient (Градиент)",
     "плавный переход между двумя цветами. Используется в нашем UI для кнопок и панелей."),
    ("Cancellation Token",
     "механизм остановки асинхронной операции. Нажав кнопку «Стоп» в чате, "
     "пользователь отменяет запрос к ИИ, не дожидаясь ответа."),
]

for term, definition in terms:
    add_term(term, definition)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  3. БЭКЕНД
# ══════════════════════════════════════════════════════════════════

add_heading("3. БЭКЕНД (BACKEND) — СЕРВЕРНАЯ И ЛОГИЧЕСКАЯ ЧАСТЬ", 1)

add_body(
    "Бэкенд — это «мозг» приложения. Он не виден пользователю, но именно здесь происходит "
    "вся реальная работа: отправка запросов к ИИ, сохранение данных, принятие решений о "
    "маршрутизации. В нашем проекте бэкенд разделён на три части: модели данных, "
    "работа с базой данных и сервисы.", indent=True
)

# ── 3.1 Язык и платформа
add_heading("3.1 Язык программирования и платформа", 2)

add_body("Технологический стек (набор технологий) нашего проекта:")

stack = [
    ("C# .NET 10",
     "основной язык. Современный, быстрый, строго типизированный язык Microsoft. "
     "«Строго типизированный» означает: каждая переменная имеет чёткий тип "
     "(строка, число, объект), что снижает количество ошибок."),
    ("Windows Forms",
     "библиотека для создания оконного интерфейса. Используется как фронтенд-движок."),
    ("SQLite (Microsoft.Data.Sqlite)",
     "встроенная база данных. Весь файл БД = один файл focusflow.db на диске."),
    ("Newtonsoft.Json",
     "библиотека для работы с JSON: сериализация конфига, десериализация ответов от ИИ."),
    ("HttpClient (.NET)",
     "встроенный HTTP-клиент для отправки запросов к API провайдеров ИИ."),
]
for name, desc in stack:
    add_term(name, desc)

# ── 3.2 Модели данных
add_heading("3.2 Модели данных (папка Models/)", 2)

add_body(
    "Модели данных — это классы C#, которые описывают объекты нашего приложения. "
    "Это как «схема базы данных», но на языке программирования. "
    "Каждый класс соответствует таблице в базе данных.", indent=True
)

models_info = [
    ("AiMessage.cs — сообщение в чате",
     [
         "int Id — уникальный номер сообщения (первичный ключ в БД)",
         "int ConversationId — к какому чату относится (внешний ключ)",
         "MessageRole Role — кто написал: User (пользователь), Assistant (ИИ), System (инструкция)",
         "string Content — текст сообщения",
         "DateTime CreatedAt — дата и время создания",
         "int TokensUsed — сколько токенов потрачено на это сообщение",
         "string? ModelUsed — какая модель ИИ написала ответ",
         "bool IsError — была ли ошибка при получении ответа",
     ]),
    ("Conversation.cs — чат (диалог)",
     [
         "int Id — уникальный номер чата",
         "string Title — название чата (автоматически генерируется ИИ)",
         "string AgentId — какой агент используется в этом чате",
         "string Model — модель ИИ по умолчанию для этого чата",
         "bool IsPinned — закреплён ли чат в боковой панели",
         "DateTime CreatedAt — когда создан",
         "DateTime UpdatedAt — когда последний раз обновлялся",
     ]),
    ("AIAgent.cs — ИИ-агент",
     [
         "string Id — уникальный текстовый идентификатор (например: 'code', 'writer')",
         "string Name — отображаемое имя ('Code Master', 'Content Writer')",
         "string Description — краткое описание назначения агента",
         "string SystemPrompt — базовые инструкции для агента (его «личность»)",
         "string Emoji — иконка агента (💻, ✍️, 📊...)",
         "string ColorHex — цвет в формате #RRGGBB",
         "string Model — предпочтительная модель ИИ для этого агента",
         "float Temperature — «творческость» от 0.0 (точный) до 1.0 (креативный)",
         "int MaxTokens — максимальная длина ответа",
     ]),
    ("AppConfig.cs — конфигурация приложения",
     [
         "string AnthropicApiKey — ключ для Claude (Anthropic)",
         "string OpenAiApiKey — ключ для GPT (OpenAI)",
         "string GeminiApiKey — ключ для Gemini (Google)",
         "string MistralApiKey — ключ для Mistral",
         "string GroqApiKey — ключ для Groq (быстрые бесплатные модели)",
         "string DefaultModel — модель ИИ по умолчанию",
         "OrchestrationMode DefaultOrchestration — режим AUTO, FUSION или MANUAL",
         "float Temperature — глобальная настройка «творческости»",
         "bool AutoTitleChats — автоматически создавать название для чата",
     ]),
    ("WorkflowStep.cs — шаг воркфлоу",
     [
         "int WorkflowId — к какому воркфлоу относится",
         "int StepOrder — порядковый номер шага в цепочке",
         "string AgentId — какой агент выполняет этот шаг",
         "string StepName — название шага",
         "string Instruction — дополнительная инструкция для этого шага",
     ]),
]

for model_name, fields in models_info:
    add_heading(f"Класс {model_name}", 3)
    for field in fields:
        p = doc.add_paragraph(f"• {field}", style='List Bullet')
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Courier New'

# ── 3.3 База данных
add_heading("3.3 База данных и репозитории (папка Data/)", 2)

add_body(
    "SQLite-база данных хранится в файле focusflow.db по пути "
    "%AppData%\\FocusFlowAI\\focusflow.db. "
    "При первом запуске программа автоматически создаёт все таблицы и заполняет их "
    "встроенными агентами (seeding).", indent=True
)

add_heading("Структура базы данных:", 3)
add_body("Таблица Conversations (чаты):")
add_code(
    "CREATE TABLE Conversations (\n"
    "    Id        INTEGER PRIMARY KEY AUTOINCREMENT,  -- автонумерация\n"
    "    Title     TEXT    NOT NULL,                   -- название чата\n"
    "    AgentId   TEXT    NOT NULL,                   -- ID агента\n"
    "    Model     TEXT    NOT NULL,                   -- модель ИИ\n"
    "    IsPinned  INTEGER NOT NULL DEFAULT 0,         -- 0=нет, 1=да\n"
    "    CreatedAt TEXT    NOT NULL,                   -- дата создания\n"
    "    UpdatedAt TEXT    NOT NULL                    -- дата обновления\n"
    ");"
)

add_body("Таблица Messages (сообщения):")
add_code(
    "CREATE TABLE Messages (\n"
    "    Id             INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "    ConversationId INTEGER NOT NULL,  -- связь с таблицей Conversations\n"
    "    Role           TEXT    NOT NULL,  -- 'User', 'Assistant', 'System'\n"
    "    Content        TEXT    NOT NULL,  -- текст сообщения\n"
    "    CreatedAt      TEXT    NOT NULL,\n"
    "    TokensUsed     INTEGER NOT NULL DEFAULT 0,\n"
    "    ModelUsed      TEXT,              -- NULL если сообщение от пользователя\n"
    "    IsError        INTEGER NOT NULL DEFAULT 0,\n"
    "    FOREIGN KEY (ConversationId) REFERENCES Conversations(Id) ON DELETE CASCADE\n"
    ");"
)

add_body(
    "FOREIGN KEY — внешний ключ — связывает таблицы: каждое сообщение "
    "знает, к какому чату оно относится. ON DELETE CASCADE означает: "
    "если удалить чат — все его сообщения удалятся автоматически."
)

add_heading("Репозитории:", 3)
repos = [
    ("ConversationRepository",
     "Методы: GetAll() — получить все чаты; GetById(id) — найти чат по ID; "
     "Create() — создать новый чат; UpdateTitle() — переименовать; "
     "TogglePin() — закрепить/открепить; Delete() — удалить."),
    ("MessageRepository",
     "Методы: GetByConversation(id) — все сообщения чата; "
     "Add(message) — добавить сообщение; GetLastN(id, n) — последние N сообщений."),
    ("AgentRepository",
     "Методы: GetAll() — все агенты; GetById(id) — найти агента; "
     "Upsert(agent) — создать или обновить; Delete(id) — удалить."),
    ("WorkflowRepository",
     "Методы: GetAll() — все воркфлоу; Save(workflow) — сохранить; "
     "Delete(id) — удалить."),
]
for name, desc in repos:
    add_term(name, desc)

add_page_break()

# ── 3.4 Сервисы
add_heading("3.4 Сервисы и бизнес-логика (папка Services/)", 2)

add_body(
    "Сервисы — это «рабочие лошадки» бэкенда. Они содержат всю бизнес-логику: "
    "как отправить запрос ИИ, как выбрать лучшую модель, как запустить цепочку агентов. "
    "Сервисы ничего не знают о кнопках и окнах — только о логике.", indent=True
)

# ── 3.5 Провайдеры
add_heading("3.5 Провайдеры ИИ", 2)

add_body(
    "Провайдер — класс, который умеет отправлять запросы к конкретному ИИ-сервису. "
    "Все провайдеры реализуют единый интерфейс ILLMProvider:", indent=True
)

add_code(
    "public interface ILLMProvider\n"
    "{\n"
    "    ProviderType ProviderType { get; }  // Anthropic, OpenAI, Google...\n"
    "    bool IsConfigured { get; }          // настроен ли API-ключ?\n"
    "    Task<AIResponse> SendAsync(         // отправить запрос\n"
    "        string systemPrompt,            // инструкция агента\n"
    "        List<ChatMessage> history,      // история диалога\n"
    "        string userMessage,             // вопрос пользователя\n"
    "        string modelApiName,            // ID модели ('claude-opus-4-6')\n"
    "        float temperature,              // 0.0–1.0\n"
    "        int maxTokens,                  // макс. длина ответа\n"
    "        CancellationToken ct            // токен отмены\n"
    "    );\n"
    "}"
)

providers_list = [
    ("AnthropicProvider",
     "Отправляет POST-запрос на https://api.anthropic.com/v1/messages. "
     "Обслуживает модели: claude-opus-4-6, claude-sonnet-4-6, claude-haiku-4-5."),
    ("GeminiProvider",
     "Отправляет запрос в Google AI Studio API. "
     "Обслуживает модели: gemini-2.5-pro, gemini-2.0-flash."),
    ("OpenAICompatProvider",
     "Универсальный провайдер, работающий с любым API в OpenAI-совместимом формате. "
     "Используется для: OpenAI (GPT-4o), Mistral, Groq. "
     "Groq примечателен скоростью — отвечает быстрее всех."),
]
for name, desc in providers_list:
    add_term(name, desc)

add_body("Класс AIResponse — что возвращает провайдер:")
add_code(
    "public class AIResponse\n"
    "{\n"
    "    string  Text;         // текст ответа ИИ\n"
    "    int     InputTokens;  // токены запроса (сколько мы отправили)\n"
    "    int     OutputTokens; // токены ответа (сколько ИИ написал)\n"
    "    string? Model;        // какая модель ответила\n"
    "    string? Provider;     // название провайдера\n"
    "    bool    Success;      // true = всё хорошо, false = ошибка\n"
    "    string? Error;        // текст ошибки (если Success = false)\n"
    "}"
)

# ── 3.6 AIRouter
add_heading("3.6 Маршрутизатор запросов (AIRouter)", 2)

add_body(
    "AIRouter — интеллектуальный компонент, который анализирует запрос пользователя и "
    "решает, какой ИИ лучше справится. Это ключевая особенность FocusFlow AI.", indent=True
)

add_heading("Как работает классификация запроса:", 3)
add_body(
    "Маршрутизатор проверяет текст запроса на наличие ключевых слов. "
    "Например, если запрос содержит слова «код», «баг», «функция», «python» — "
    "тип запроса определяется как Code (программирование)."
)

types_table = [
    ("Code (программирование)",
     "код, code, баг, bug, функция, python, c#, java, debug, алгоритм",
     "Лучший выбор: Code-специализированные модели с высоким QualityScore"),
    ("Math (математика)",
     "посчитай, вычисли, уравнение, формула, вероятность, интеграл",
     "Лучший выбор: модели с флагом IsMath = true"),
    ("Creative (творческий)",
     "напиши рассказ, стихотворение, статья, копирайтинг, слоган",
     "Лучший выбор: модели с флагом IsCreative = true"),
    ("Analysis (анализ)",
     "проанализируй, сравни, объясни, плюсы и минусы, стратегия",
     "Лучший выбор: модели с флагом IsAnalysis = true"),
    ("Simple (простой)",
     "короткий запрос (< 80 символов), нет специальных ключевых слов",
     "Лучший выбор: самая быстрая и дешёвая доступная модель"),
]

for qtype, keywords, choice in types_table:
    add_heading(f"Тип: {qtype}", 3)
    add_body(f"Ключевые слова: {keywords}")
    add_body(f"Стратегия выбора: {choice}")

add_heading("Система оценки моделей (Scoring):", 3)
add_body(
    "Каждая доступная модель получает числовой балл. "
    "Побеждает модель с наивысшим баллом."
)
add_code(
    "score = QualityScore * 10             // базовый балл качества\n"
    "score += (IsCodeExpert ? 30 : 0)      // бонус для кода\n"
    "score += (IsMath ? 30 : 0)            // бонус для математики\n"
    "score += (IsCreative ? 25 : 0)        // бонус для творческих задач\n"
    "score -= CostPer1K * 5000             // штраф за высокую стоимость (для простых запросов)"
)

# ── 3.7 OrchestrationService
add_heading("3.7 Оркестратор (OrchestrationService)", 2)

add_body(
    "OrchestrationService — главный дирижёр. Он получает запрос пользователя, "
    "узнаёт у маршрутизатора, какие модели использовать, и управляет всем процессом "
    "получения ответа. Поддерживает три режима работы:", indent=True
)

modes = [
    ("AUTO-режим",
     "Маршрутизатор выбирает одну лучшую модель → запрос отправляется в неё → "
     "ответ возвращается пользователю. Самый быстрый и экономный режим."),
    ("FUSION-режим",
     "Маршрутизатор выбирает 2 модели от разных провайдеров → "
     "запросы отправляются ПАРАЛЛЕЛЬНО (Task.WhenAll) → "
     "оба ответа передаются третьей модели (синтезатору) → "
     "синтезатор объединяет лучшее из обоих → финальный ответ возвращается пользователю. "
     "Качество выше, но дороже и медленнее."),
    ("MANUAL-режим",
     "Пользователь сам выбирает провайдера и модель из выпадающих списков. "
     "Маршрутизатор не вмешивается."),
]
for mode_name, desc in modes:
    add_heading(f"Режим {mode_name}:", 3)
    add_body(desc)

add_body("Промпт синтезатора (инструкция для объединения ответов):")
add_code(
    "\"Ты — синтезатор ответов. Тебе даны два ответа от разных AI моделей.\n"
    " Твоя задача: создать ОДИН идеальный ответ, объединив лучшее из обоих.\n"
    " - Возьми сильные стороны каждого ответа\n"
    " - Убери дубликаты и противоречия\n"
    " - Ответ должен быть лучше любого из исходных\""
)

# ── 3.8 WorkflowService
add_heading("3.8 Воркфлоу — цепочки агентов (WorkflowService)", 2)

add_body(
    "WorkflowService выполняет автоматические цепочки задач. "
    "Пользователь создаёт воркфлоу из нескольких шагов, каждый шаг выполняет свой агент.", indent=True
)

add_body("Логика выполнения:")
add_code(
    "Ввод пользователя → Шаг 1 (Агент A обрабатывает) → выход Шага 1\n"
    "→ Шаг 2 (Агент B обрабатывает выход Шага 1) → выход Шага 2\n"
    "→ Шаг 3 (Агент C обрабатывает выход Шага 2) → финальный результат"
)

add_body("Пример реального воркфлоу:")
steps = [
    "Шаг 1 — Research Pro: «Исследуй тему X, дай структурированный обзор»",
    "Шаг 2 — Content Writer: «Перепиши исследование в виде статьи»",
    "Шаг 3 — Marketing Expert: «Добавь заголовок и призыв к действию»",
]
for step in steps:
    p = doc.add_paragraph(f"• {step}")
    p.runs[0].font.size = Pt(11)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  4. ФРОНТЕНД
# ══════════════════════════════════════════════════════════════════

add_heading("4. ФРОНТЕНД (FRONTEND) — ВИЗУАЛЬНЫЙ ИНТЕРФЕЙС", 1)

add_body(
    "Фронтенд отвечает за всё, что видит пользователь: окна, кнопки, цвета, шрифты, анимации. "
    "В нашем проекте он построен на Windows Forms с кастомными (самописными) компонентами "
    "и единой дизайн-системой.", indent=True
)

# ── 4.1 Дизайн-система
add_heading("4.1 Дизайн-система (Theme.cs)", 2)

add_body(
    "Theme.cs — центральный файл всех визуальных констант. "
    "Вместо того чтобы прописывать цвет кнопки в каждом месте, "
    "мы пишем Theme.Primary — и цвет берётся из одного места. "
    "Если нужно поменять цвет — меняем только в Theme.cs.", indent=True
)

add_heading("Цветовая палитра:", 3)
colors = [
    ("Background #0D0D12", "почти чёрный — основной фон приложения"),
    ("SidebarBg #111118", "чуть светлее — фон боковой панели"),
    ("CardBg #1A1A26", "тёмно-фиолетовый — карточки и панели"),
    ("Primary #7C5CFC", "фиолетовый — основной акцентный цвет кнопок"),
    ("PrimaryLight #9B7FFF", "светло-фиолетовый — ховер-состояние кнопок"),
    ("TextPrimary #E8E8F0", "почти белый — основной текст"),
    ("TextSecondary #8C8CAF", "серо-фиолетовый — вторичный текст"),
    ("Success #48C774", "зелёный — успешные операции"),
    ("Error #FF5050", "красный — ошибки"),
    ("Warning #FFB43C", "оранжевый — предупреждения"),
]
for color_name, desc in colors:
    add_term(color_name, desc)

add_heading("Шрифты:", 3)
fonts = [
    ("FontH1 — Segoe UI 20pt Bold", "заголовки страниц"),
    ("FontH2 — Segoe UI 15pt Bold", "заголовки секций"),
    ("FontH3 — Segoe UI 12pt Bold", "заголовки карточек"),
    ("FontBody — Segoe UI 10pt", "основной текст"),
    ("FontMono — Cascadia Code 9.5pt", "код в сообщениях чата"),
]
for font_name, usage in fonts:
    add_term(font_name, usage)

add_heading("Вспомогательные методы Theme:", 3)
methods = [
    ("DrawRoundedRect()", "рисует прямоугольник со скруглёнными углами (используется везде)"),
    ("DrawGradientRect()", "рисует прямоугольник с градиентной заливкой (кнопки, панели)"),
    ("Darken() / Lighten()", "затемнить/осветлить цвет на N% (для hover-эффектов)"),
    ("WithAlpha()", "добавить прозрачность к цвету"),
]
for m, d in methods:
    add_term(m, d)

# ── 4.2 Переиспользуемые компоненты
add_heading("4.2 Переиспользуемые компоненты (папка Controls/)", 2)

add_body(
    "Кастомные контролы — это компоненты, которых нет в стандартном Windows Forms "
    "и которые мы написали сами для нужного визуального стиля.", indent=True
)

controls = [
    ("RoundedPanel",
     "Панель со скруглёнными углами и опциональной тенью. "
     "Стандартный Panel в WinForms — прямоугольный и скучный. "
     "Наш RoundedPanel рисует себя через Graphics.FillPath с rounded path."),
    ("FlatButton",
     "Кнопка в стиле приложения — без системного вида Windows, "
     "с кастомным фоном, скруглёнными углами, hover-эффектом и поддержкой градиента. "
     "Переопределяет метод OnPaint для полностью кастомной отрисовки."),
    ("IconButton",
     "Маленькая кнопка только с иконкой (без текста). "
     "Используется для кнопки «Стоп» в чате и других управляющих кнопок."),
    ("MessageBubble",
     "Пузырёк сообщения в чате — как в WhatsApp/Telegram. "
     "Сообщения пользователя — справа, фиолетовые. "
     "Сообщения ИИ — слева, тёмно-серые. "
     "Содержит: текст сообщения, время, имя модели ИИ, кнопку копирования."),
    ("AgentCard",
     "Карточка агента на странице «Агенты». "
     "Показывает: иконку-эмодзи, имя, описание, цветную полоску. "
     "При клике — открывает этого агента в чате."),
    ("InputDialog",
     "Модальное окно для ввода текста. Используется при переименовании чата "
     "и создании нового агента."),
]
for name, desc in controls:
    add_heading(f"Компонент {name}:", 3)
    add_body(desc)

# ── 4.3 MainForm
add_heading("4.3 Главное окно (MainForm.cs)", 2)

add_body(
    "MainForm — это «скелет» всего приложения. Он содержит боковую панель (sidebar) "
    "с навигацией и главную область контента, куда загружаются страницы. "
    "Также MainForm хранит все сервисы и репозитории как статические поля — "
    "это «точка входа» для всей бизнес-логики.", indent=True
)

add_body("Структура главного окна:")
add_code(
    "┌─────────────────────────────────────────────────────┐\n"
    "│  SIDEBAR (250px)     │  CONTENT AREA               │\n"
    "│  ─────────────────   │  ─────────────────────────  │\n"
    "│  ✨ FocusFlow AI     │                             │\n"
    "│  [+ Новый чат]       │  Здесь отображается         │\n"
    "│  🔍 Поиск...         │  текущая активная страница: │\n"
    "│  ─────────────────   │  ChatPage / AgentsPage /    │\n"
    "│  Список чатов        │  AutomationPage /           │\n"
    "│  ─────────────────   │  HistoryPage /              │\n"
    "│  💬 Чат              │  SettingsPage               │\n"
    "│  🤖 Агенты           │                             │\n"
    "│  ⚡ Автоматизация    │                             │\n"
    "│  📋 История          │                             │\n"
    "│  ⚙️  Настройки        │                             │\n"
    "└─────────────────────────────────────────────────────┘"
)

add_body("Статические сервисы MainForm:")
add_code(
    "MainForm.Config       — AppConfig (настройки и ключи)\n"
    "MainForm.AI           — AIService (фасад для отправки запросов)\n"
    "MainForm.Orchestrator — OrchestrationService (главный дирижёр)\n"
    "MainForm.ConvRepo     — ConversationRepository (чаты)\n"
    "MainForm.MsgRepo      — MessageRepository (сообщения)\n"
    "MainForm.AgentRepo    — AgentRepository (агенты)\n"
    "MainForm.WfRepo       — WorkflowRepository (воркфлоу)"
)

# ── 4.4 ChatPage
add_heading("4.4 Страница чата (ChatPage.cs)", 2)

add_body(
    "ChatPage — главная страница приложения. Здесь происходит весь диалог с ИИ. "
    "Это UserControl (переиспользуемый визуальный компонент), который загружается "
    "в центральную область MainForm.", indent=True
)

add_body("Визуальная структура ChatPage:")
add_code(
    "┌───────────────────────────────────────────────┐\n"
    "│ TOP BAR                                       │\n"
    "│  Название чата        [Выбор агента ▼]        │\n"
    "│  Статус модели        [AUTO][FUSION][ВРУЧНУЮ] │\n"
    "│  (ручной режим: [Провайдер ▼] [Модель ▼])    │\n"
    "├───────────────────────────────────────────────┤\n"
    "│ THINKING BAR (скрыт по умолчанию)             │\n"
    "│  ⚡ Claude думает...                          │\n"
    "├───────────────────────────────────────────────┤\n"
    "│                                               │\n"
    "│  ОБЛАСТЬ СООБЩЕНИЙ (прокручиваемая)           │\n"
    "│                                               │\n"
    "│       ┌─────────────────────┐                 │\n"
    "│       │ Сообщение от ИИ     │                 │\n"
    "│       └─────────────────────┘                 │\n"
    "│  ┌─────────────────────┐                     │\n"
    "│  │ Ваш вопрос          │                     │\n"
    "│  └─────────────────────┘                     │\n"
    "│                                               │\n"
    "├───────────────────────────────────────────────┤\n"
    "│ INPUT AREA                                    │\n"
    "│  [Введите сообщение...       ] [▶ Отправить]  │\n"
    "└───────────────────────────────────────────────┘"
)

add_heading("Процесс отправки сообщения (события):", 3)
steps_chat = [
    "Пользователь вводит текст в RichTextBox (_inputBox)",
    "Нажимает Enter или кнопку «Отправить» (_sendBtn)",
    "Вызывается метод SendMessageAsync()",
    "Создаётся CancellationTokenSource (для кнопки «Стоп»)",
    "Сообщение пользователя сохраняется в БД (MessageRepository.Add)",
    "Отображается «Thinking bar» с анимацией точек",
    "Вызывается MainForm.Orchestrator.RunAsync() — запрос к ИИ",
    "Получен ответ → создаётся MessageBubble → добавляется в _messagesArea",
    "Ответ сохраняется в БД",
    "Если AutoTitleChats = true — ИИ генерирует название чата",
]
for i, step in enumerate(steps_chat, 1):
    p = doc.add_paragraph(f"{i}. {step}")
    p.runs[0].font.size = Pt(11)

# ── 4.5 AgentsPage
add_heading("4.5 Страница агентов (AgentsPage.cs)", 2)

add_body(
    "Страница для управления ИИ-агентами. Отображает сетку карточек (AgentCard) "
    "для каждого агента. Встроенные агенты нельзя удалить (IsBuiltIn = true), "
    "пользовательские — можно редактировать и удалять.", indent=True
)

agents_list = [
    ("✨ FocusFlow AI", "универсальный помощник"),
    ("💻 Code Master", "эксперт по программированию"),
    ("✍️ Content Writer", "копирайтер и контент-стратег"),
    ("📊 Data Analyst", "аналитик данных и бизнес-аналитика"),
    ("📣 Marketing Expert", "маркетолог и стратег"),
    ("🔬 Research Pro", "глубокое исследование"),
    ("🎨 Design Advisor", "UI/UX дизайн"),
    ("🎓 AI Teacher", "объяснение сложных тем"),
]
for agent, desc in agents_list:
    add_term(agent, desc)

# ── 4.6 AutomationPage
add_heading("4.6 Страница автоматизации (AutomationPage.cs)", 2)

add_body(
    "Здесь пользователь создаёт и запускает воркфлоу (цепочки агентов). "
    "Интерфейс позволяет: добавлять шаги, выбирать агента для каждого шага, "
    "задавать инструкцию, запускать воркфлоу с начальным текстом и наблюдать "
    "за прогрессом выполнения.", indent=True
)

# ── 4.7 HistoryPage
add_heading("4.7 История чатов (HistoryPage.cs)", 2)

add_body(
    "Показывает все сохранённые разговоры с датой, количеством сообщений "
    "и последним сообщением. Позволяет быстро открыть любой старый чат "
    "и продолжить диалог.", indent=True
)

# ── 4.8 SettingsPage
add_heading("4.8 Настройки (SettingsPage.cs)", 2)

add_body(
    "Страница настроек позволяет вводить API-ключи всех провайдеров. "
    "После сохранения вызывается MainForm.ReloadServices() — все сервисы "
    "пересоздаются с новыми ключами. Ключи сохраняются в config.json "
    "по пути %AppData%\\FocusFlowAI\\config.json.", indent=True
)

settings_fields = [
    "API ключ Anthropic (Claude) — обязательный для базовой работы",
    "API ключ OpenAI (GPT-4o, GPT-4o mini)",
    "API ключ Google (Gemini 2.5 Pro, Gemini 2.0 Flash)",
    "API ключ Mistral (Mistral Large, Mistral Small)",
    "API ключ Groq (Llama 3.3 70B — бесплатный и быстрый)",
    "Модель по умолчанию",
    "Режим оркестрации по умолчанию (AUTO / FUSION / MANUAL)",
    "Temperature (творческость) — ползунок 0.0 – 1.0",
    "Макс. токенов в ответе",
    "Автоматическое создание названий чатов",
]
for field in settings_fields:
    p = doc.add_paragraph(f"• {field}")
    p.runs[0].font.size = Pt(11)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  5. ПОЛНЫЙ ЦИКЛ ЗАПРОСА
# ══════════════════════════════════════════════════════════════════

add_heading("5. КАК ВСЁ РАБОТАЕТ ВМЕСТЕ: ПОЛНЫЙ ЦИКЛ ЗАПРОСА", 1)

add_body(
    "Рассмотрим, что происходит в программе от момента нажатия «Отправить» "
    "до появления ответа на экране.", indent=True
)

add_code(
    "ПОЛЬЗОВАТЕЛЬ нажимает Enter\n"
    "         │\n"
    "         ▼\n"
    "ChatPage.SendMessageAsync()\n"
    "  ├─ Сохранить сообщение пользователя → MessageRepository.Add()\n"
    "  ├─ Показать ThinkingBar («ИИ думает...»)\n"
    "  └─ Вызвать OrchestrationService.RunAsync()\n"
    "              │\n"
    "              ▼\n"
    "         AIRouter.Route(message, mode)\n"
    "           ├─ Classify(message) → определить тип запроса\n"
    "           ├─ GetAvailableModels() → список настроенных моделей\n"
    "           └─ PickBest() → выбрать лучшую модель (Score-алгоритм)\n"
    "              │\n"
    "              ▼ (режим AUTO)\n"
    "         ProviderRegistry.SendAsync(model, ...)\n"
    "           └─ AnthropicProvider.SendAsync()   (или OpenAI / Gemini...)\n"
    "                  └─ HttpClient.PostAsync(API_URL, json_body)\n"
    "                  └─ Ждём ответ от сервера ИИ...\n"
    "                  └─ Десериализовать JSON → AIResponse\n"
    "              │\n"
    "              ▼\n"
    "         OrchestrationResult (Success=true, Text=ответ)\n"
    "              │\n"
    "              ▼\n"
    "ChatPage: скрыть ThinkingBar\n"
    "  ├─ Создать MessageBubble(ответ)\n"
    "  ├─ Добавить в _messagesArea\n"
    "  ├─ Сохранить ответ → MessageRepository.Add()\n"
    "  └─ Если AutoTitle → GenerateTitleAsync() → ConvRepo.UpdateTitle()"
)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  6. СТРУКТУРА ФАЙЛОВ
# ══════════════════════════════════════════════════════════════════

add_heading("6. СТРУКТУРА ФАЙЛОВ ПРОЕКТА", 1)

add_code(
    "FocusFlow LMS/\n"
    "├── Program.cs              ← точка входа, запуск приложения\n"
    "├── Theme.cs                ← дизайн-система: цвета, шрифты, методы рисования\n"
    "│\n"
    "├── Models/                 ← БЭКЕНД: модели данных\n"
    "│   ├── AiMessage.cs        ← сообщение в чате\n"
    "│   ├── Conversation.cs     ← чат/диалог\n"
    "│   ├── AIAgent.cs          ← агент ИИ\n"
    "│   ├── AppConfig.cs        ← конфигурация + сохранение\n"
    "│   ├── AIProvider.cs       ← перечисления: ProviderType, QueryType, ModelInfo\n"
    "│   └── WorkflowStep.cs     ← шаг воркфлоу\n"
    "│\n"
    "├── Data/                   ← БЭКЕНД: работа с базой данных\n"
    "│   ├── DatabaseManager.cs  ← создание таблиц, seed-данные\n"
    "│   ├── ConversationRepository.cs\n"
    "│   ├── MessageRepository.cs\n"
    "│   ├── AgentRepository.cs\n"
    "│   └── WorkflowRepository.cs\n"
    "│\n"
    "├── Services/               ← БЭКЕНД: бизнес-логика\n"
    "│   ├── ILLMProvider.cs     ← интерфейс провайдера ИИ\n"
    "│   ├── AnthropicProvider.cs← Claude API\n"
    "│   ├── GeminiProvider.cs   ← Gemini API\n"
    "│   ├── OpenAICompatProvider.cs ← OpenAI / Mistral / Groq\n"
    "│   ├── ProviderRegistry.cs ← реестр всех провайдеров\n"
    "│   ├── AIRouter.cs         ← маршрутизатор (выбор модели)\n"
    "│   ├── OrchestrationService.cs ← AUTO / FUSION / MANUAL\n"
    "│   ├── AIService.cs        ← фасад для упрощённого вызова\n"
    "│   └── WorkflowService.cs  ← цепочки агентов\n"
    "│\n"
    "├── Controls/               ← ФРОНТЕНД: кастомные компоненты\n"
    "│   ├── RoundedPanel.cs\n"
    "│   ├── FlatButton.cs\n"
    "│   ├── IconButton.cs\n"
    "│   ├── MessageBubble.cs\n"
    "│   ├── AgentCard.cs\n"
    "│   └── InputDialog.cs\n"
    "│\n"
    "└── Forms/                  ← ФРОНТЕНД: страницы и главное окно\n"
    "    ├── MainForm.cs         ← главное окно с боковой панелью\n"
    "    ├── ChatPage.cs         ← страница чата\n"
    "    ├── AgentsPage.cs       ← управление агентами\n"
    "    ├── AutomationPage.cs   ← воркфлоу\n"
    "    ├── HistoryPage.cs      ← история чатов\n"
    "    └── SettingsPage.cs     ← настройки и API-ключи"
)

add_page_break()

# ══════════════════════════════════════════════════════════════════
#  7. ЗАКЛЮЧЕНИЕ
# ══════════════════════════════════════════════════════════════════

add_heading("7. ЗАКЛЮЧЕНИЕ", 1)

add_body(
    "FocusFlow AI представляет собой полноценное настольное приложение, "
    "демонстрирующее современные подходы к интеграции больших языковых моделей. "
    "Проект реализует три ключевые идеи: унификацию (один интерфейс для пяти провайдеров), "
    "интеллектуальную маршрутизацию (AUTO-режим) и синергию моделей (FUSION-режим).", indent=True
)

add_body(
    "С точки зрения архитектуры проект строго разделён на бэкенд (Models, Data, Services) "
    "и фронтенд (Controls, Forms, Theme). Бэкенд не знает о кнопках и окнах, "
    "фронтенд не знает о SQL и HTTP — это называется принципом разделения ответственности "
    "(Separation of Concerns), одним из ключевых принципов хорошей архитектуры ПО.", indent=True
)

add_body(
    "Воркфлоу-система реализует концепцию мультиагентной архитектуры: "
    "каждый агент специализируется на своей задаче, а их последовательное взаимодействие "
    "позволяет решать сложные многоэтапные задачи автоматически.", indent=True
)

add_heading("Использованные технологии:", 2)
final_tech = [
    "C# .NET 10 — язык программирования и платформа",
    "Windows Forms — фреймворк для создания оконного интерфейса",
    "SQLite + Microsoft.Data.Sqlite — локальная база данных",
    "Newtonsoft.Json — сериализация/десериализация JSON",
    "HttpClient (.NET) — HTTP-запросы к API провайдеров",
    "Anthropic Claude API — основной AI-провайдер",
    "OpenAI API — провайдер GPT-моделей",
    "Google Gemini API — провайдер Gemini-моделей",
    "Mistral AI API — быстрые европейские модели",
    "Groq API — сверхбыстрый инференс open-source моделей",
]
for tech in final_tech:
    p = doc.add_paragraph(f"• {tech}")
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'

# ── Сохранение ──────────────────────────────────────────────────
output_path = r"C:\Users\NOXQD\source\repos\FocusFlow LMS\FocusFlow_AI_Coursework.docx"
doc.save(output_path)
print(f"OK! Document saved: {output_path}")

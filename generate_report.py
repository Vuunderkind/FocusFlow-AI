# -*- coding: utf-8 -*-
"""
Генератор отчёта по учебной практике — FocusFlow AI
Структура соответствует методическим рекомендациям (ПМ3, Колледж Хекслет)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────
#  Страница: поля A4
# ─────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ─────────────────────────────────────────────────────────────────
#  Вспомогательные функции
# ─────────────────────────────────────────────────────────────────
def set_run_font(run, size=14, bold=False, italic=False, color=None):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # CyrillicFont fix
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),   'Times New Roman')
    rFonts.set(qn('w:cs'),      'Times New Roman')
    rPr.insert(0, rFonts)


def para(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14,
         bold=False, italic=False, indent=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(text, level=1, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def heading_left(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def bullet(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent       = Cm(1.25)
    pf.first_line_indent = Cm(-0.5)
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run('– ' + text)
    set_run_font(run, size=size)
    return p


def page_break():
    doc.add_page_break()


def add_table_row(table, cells_data, bold_row=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after  = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=11, bold=bold_row)
    return row


def add_code_block(text):
    """Добавляет блок кода с моноширинным шрифтом"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # background shade
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F5F5F5')
    p._p.get_or_add_pPr().append(shd)
    return p


# ═════════════════════════════════════════════════════════════════
#  ТИТУЛЬНЫЙ ЛИСТ
# ═════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(0)
pf.space_after  = Pt(6)
run = p.add_run('ТОО «Колледж Хекслет»')
set_run_font(run, size=14, bold=True)

# Пустые строки
for _ in range(6):
    para('', space_before=0, space_after=0)

heading('ОТЧЕТ ПО УЧЕБНОЙ ПРАКТИКЕ', size=16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf2 = p2.paragraph_format
pf2.space_before = Pt(6)
pf2.space_after  = Pt(4)
r1 = p2.add_run('на тему: ')
set_run_font(r1, size=14, bold=True)
r2 = p2.add_run('«Разработка настольного приложения "FocusFlow AI" (Windows Forms)»')
set_run_font(r2, size=14, bold=True)

para('', space_before=0, space_after=4)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf3 = p3.paragraph_format
pf3.space_before = Pt(4)
pf3.space_after  = Pt(4)
r = p3.add_run('по модулю: ПМ3 «Разработка модулей программного обеспечения для компьютерных систем»')
set_run_font(r, size=14, bold=True)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf4 = p4.paragraph_format
pf4.space_before = Pt(4)
pf4.space_after  = Pt(4)
r = p4.add_run('Специальность: 06130100 – Программное обеспечение (по видам)')
set_run_font(r, size=14, bold=True)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf5 = p5.paragraph_format
pf5.space_before = Pt(4)
pf5.space_after  = Pt(4)
r = p5.add_run('Квалификация: 4S06130105 – Техник информационных систем')
set_run_font(r, size=14, bold=True)

for _ in range(4):
    para('', space_before=0, space_after=0)

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf6 = p6.paragraph_format
pf6.space_before = Pt(0)
pf6.space_after  = Pt(4)
r = p6.add_run('Выполнил(-а): студент 3 курса')
set_run_font(r, size=14, bold=True)

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf7 = p7.paragraph_format
pf7.space_before = Pt(0)
pf7.space_after  = Pt(4)
r = p7.add_run('группы 22 ТИС')
set_run_font(r, size=14, bold=True)

p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf8 = p8.paragraph_format
pf8.space_before = Pt(0)
pf8.space_after  = Pt(0)
r = p8.add_run('Vuunderkind')
set_run_font(r, size=14, bold=True)

for _ in range(3):
    para('', space_before=0, space_after=0)

# Допуск и оценка в одну строку
tbl_title = doc.add_table(rows=1, cols=2)
tbl_title.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_title.style = 'Table Grid'
# Убрать рамку
from docx.oxml.ns import nsmap
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
remove_table_borders(tbl_title)

c0 = tbl_title.rows[0].cells[0]
c1 = tbl_title.rows[0].cells[1]

p_l = c0.paragraphs[0]
p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p_l.add_run('Допущен(-а) к защите «____»  20___г.\nПреподаватель: Бенли Р.А.\n(подпись)')
set_run_font(r, size=12, bold=True)

p_r = c1.paragraphs[0]
p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p_r.add_run('Защитил(-а) с оценкой:\n\n(буква) (циф.экв.) (балл)')
set_run_font(r, size=12, bold=True)

para('', space_before=0, space_after=6)

p_city = doc.add_paragraph()
p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_city = p_city.paragraph_format
pf_city.space_before = Pt(6)
pf_city.space_after  = Pt(0)
r = p_city.add_run('Алматы 2025')
set_run_font(r, size=14, bold=True)

page_break()

# ═════════════════════════════════════════════════════════════════
#  СОДЕРЖАНИЕ
# ═════════════════════════════════════════════════════════════════
heading('СОДЕРЖАНИЕ', size=14)

contents = [
    ('ВВЕДЕНИЕ', '3', False),
    ('1. ТЕОРЕТИЧЕСКАЯ ЧАСТЬ', '4', True),
    ('1.1 Анализ предметной области и функциональных требований', '4', False),
    ('1.2 Описание инструментов и среды разработки', '6', False),
    ('1.3 Постановка задачи (техническое задание)', '8', False),
    ('1.4 Проектирование базы данных (ER-диаграмма, 3НФ)', '10', False),
    ('1.5 Проектирование интерфейса приложения (UI/UX)', '13', False),
    ('1.6 Описание внутренней структуры приложения (архитектура)', '15', False),
    ('2. ПРАКТИЧЕСКАЯ ЧАСТЬ', '18', True),
    ('2.1 Создание проекта и подключение зависимостей', '18', False),
    ('2.2 Разработка базы данных (SQLite)', '20', False),
    ('2.3 Реализация основных функций (AI-провайдеры)', '22', False),
    ('2.4 Разработка пользовательского интерфейса', '25', False),
    ('2.5 Реализация системы оркестрации', '27', False),
    ('2.6 Обработка ошибок и валидация данных', '29', False),
    ('ЗАКЛЮЧЕНИЕ', '31', True),
    ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '32', True),
    ('ПРИЛОЖЕНИЯ', '34', True),
]

for title, page, bold in contents:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    tab_stops = p.paragraph_format.tab_stops
    # добавляем таб-стоп для номера страницы
    from docx.shared import Cm as CM
    from docx.oxml import OxmlElement as OE
    pPr = p._p.get_or_add_pPr()
    tabs = OE('w:tabs')
    tab = OE('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(int(Cm(15.5).pt * 20)))
    tabs.append(tab)
    pPr.append(tabs)
    r1 = p.add_run(title)
    set_run_font(r1, size=13, bold=bold)
    r2 = p.add_run('\t' + page)
    set_run_font(r2, size=13, bold=bold)

page_break()

# ═════════════════════════════════════════════════════════════════
#  ВВЕДЕНИЕ
# ═════════════════════════════════════════════════════════════════
heading('ВВЕДЕНИЕ', size=14)

intro_texts = [
    ('        В современном мире искусственный интеллект (ИИ) становится неотъемлемой частью '
     'профессиональной и повседневной деятельности человека. Появление мощных языковых моделей — '
     'Claude, GPT-4, Gemini, Mistral, Llama — открыло принципиально новые возможности для автоматизации '
     'интеллектуального труда. Тем не менее большинство существующих решений предоставляют доступ лишь '
     'к одному провайдеру, что ограничивает гибкость и качество генерируемых ответов.'),
    ('        Актуальность данного проекта обусловлена необходимостью создания единого настольного '
     'приложения, способного работать сразу с несколькими AI-провайдерами в рамках интеллектуальной '
     'системы оркестрации. Приложение FocusFlow AI решает задачу не просто чата с одной моделью, '
     'а целой платформы, где специализированные AI-агенты, автоматическая маршрутизация запросов '
     'и режим FUSION (синтез ответов нескольких моделей) совместно обеспечивают наилучший результат.'),
    ('        Целью данной практики является разработка настольного приложения "FocusFlow AI" на '
     'платформе .NET 10 с использованием Windows Forms, включающего:'),
]

for text in intro_texts:
    para(text)

goals = [
    'интеграцию с пятью AI-провайдерами (Anthropic, OpenAI, Google, Mistral, Groq);',
    'систему интеллектуальной маршрутизации запросов (AUTO, FUSION, MANUAL);',
    'хранение истории диалогов и агентов в локальной базе данных SQLite;',
    'механизм многошаговых воркфлоу (автоматизации) на основе цепочек агентов;',
    'современный тёмный интерфейс с поддержкой Markdown-рендеринга.',
]
for g in goals:
    bullet(g)

tasks_texts = [
    ('        Задачи практики: изучить архитектуру многопровайдерных AI-систем; спроектировать '
     'реляционную базу данных в 3НФ; реализовать асинхронные HTTP-запросы к внешним API; '
     'разработать событийно-ориентированный пользовательский интерфейс; '
     'обеспечить обработку ошибок и валидацию данных на всех уровнях приложения.'),
    ('        В результате прохождения практики были приобретены навыки работы с асинхронным '
     'программированием (async/await), интеграции REST API, проектирования SQLite-базы данных, '
     'разработки WinForms-интерфейса, а также архитектурного разделения ответственности '
     'между слоями приложения (Data, Services, Forms, Controls).'),
]
for t in tasks_texts:
    para(t)

page_break()

# ═════════════════════════════════════════════════════════════════
#  1. ТЕОРЕТИЧЕСКАЯ ЧАСТЬ
# ═════════════════════════════════════════════════════════════════
heading('1. ТЕОРЕТИЧЕСКАЯ ЧАСТЬ', size=14)

# ─── 1.1 ──────────────────────────────────────────────────────
heading_left('1.1 Анализ предметной области и функциональных требований')

para('        FocusFlow AI относится к классу десктопных AI-клиентов — программ, '
     'обеспечивающих удобный пользовательский интерфейс для взаимодействия с языковыми моделями '
     'через REST API. Предметная область включает несколько взаимосвязанных концепций.')

para('        Языковые модели большого размера (LLM) — нейросетевые системы, обученные '
     'на больших корпусах текста и способные генерировать связные, контекстуально осмысленные '
     'ответы. Провайдеры LLM предоставляют доступ к моделям через платное API с поминутной '
     'или потокенной тарификацией.', indent=1.25)

para('        Система оркестрации — программный слой, управляющий выбором модели, '
     'маршрутизацией запросов и агрегацией результатов. FocusFlow AI реализует три режима:',
     indent=1.25)

orch_items = [
    'AUTO — интеллектуальный роутер AIRouter классифицирует запрос и выбирает наиболее подходящую модель на основе ключевых слов, длины сообщения и параметров качества/стоимости каждой модели;',
    'FUSION — параллельная отправка запроса двум моделям разных провайдеров с последующим синтезом ответов третьей моделью (наиболее дешёвой из доступных);',
    'MANUAL — пользователь явно выбирает провайдера и модель из выпадающих списков.',
]
for item in orch_items:
    bullet(item)

para('        AI-агент — это набор настроек (системный промпт, модель, температура, '
     'max_tokens, эмодзи, цвет), определяющий специализацию ассистента. '
     'В приложении реализованы 8 встроенных агентов: FocusFlow AI (универсальный), '
     'Code Master, Content Writer, Data Analyst, Marketing Expert, Research Pro, '
     'Design Advisor, AI Teacher.', indent=1.25)

para('        Воркфлоу — последовательность шагов, каждый из которых выполняется '
     'определённым агентом. Выход предыдущего шага становится входом следующего, '
     'что позволяет реализовать сложные многоэтапные сценарии автоматизации.', indent=1.25)

heading_left('Функциональные требования')

func_reqs = [
    ('ФТ-01', 'Чат с AI', 'Отправка сообщений, получение ответов, сохранение истории'),
    ('ФТ-02', 'Мульти-провайдер', 'Поддержка Anthropic, OpenAI, Google, Mistral, Groq'),
    ('ФТ-03', 'Оркестрация', 'Режимы AUTO / FUSION / MANUAL'),
    ('ФТ-04', 'AI-агенты', 'Создание, редактирование, удаление агентов с системным промптом'),
    ('ФТ-05', 'Воркфлоу', 'Многошаговая автоматизация через цепочки агентов'),
    ('ФТ-06', 'История', 'Просмотр, поиск, удаление и закрепление диалогов'),
    ('ФТ-07', 'Настройки', 'Управление API-ключами, моделью по умолчанию, параметрами'),
    ('ФТ-08', 'Автозаголовок', 'Автоматическое именование диалога по первому сообщению'),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
add_table_row(tbl, ['ID', 'Функция', 'Описание'], bold_row=True)
for row_data in func_reqs:
    add_table_row(tbl, list(row_data))

# задать ширину колонок
for i, width in enumerate([Cm(2.5), Cm(4.5), Cm(9.0)]):
    for row in tbl.rows:
        row.cells[i].width = width

para('', space_before=6, space_after=0)

# ─── 1.2 ──────────────────────────────────────────────────────
heading_left('1.2 Описание инструментов и среды разработки')

tools = [
    ('Среда разработки', 'Visual Studio 2022 (Community Edition)',
     'Полнофункциональная IDE с поддержкой .NET 10, Windows Forms Designer, NuGet, отладчика и Git.'),
    ('Платформа', '.NET 10 (Windows)',
     'Последняя версия платформы, поддерживающая современные возможности C# 13, async/await, LINQ.'),
    ('Язык программирования', 'C# 13',
     'Строго типизированный объектно-ориентированный язык с поддержкой async, generics, records, pattern matching.'),
    ('UI-фреймворк', 'Windows Forms (WinForms)',
     'Событийно-ориентированный фреймворк для создания настольных приложений под Windows с богатым набором элементов управления.'),
    ('База данных', 'SQLite 3 (Microsoft.Data.Sqlite 9.x)',
     'Встраиваемая бессерверная реляционная СУБД. Файл базы данных располагается в AppData пользователя.'),
    ('HTTP-клиент', 'HttpClient (.NET)',
     'Встроенный класс для выполнения асинхронных HTTP-запросов к REST API провайдеров.'),
    ('Сериализация JSON', 'Newtonsoft.Json 13.x',
     'Библиотека для сериализации/десериализации JSON-объектов при обмене данными с API.'),
    ('Контроль версий', 'Git + GitHub',
     'Система контроля версий; репозиторий размещён на GitHub.'),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'
add_table_row(tbl2, ['Инструмент', 'Название / версия', 'Назначение'], bold_row=True)
for row_data in tools:
    add_table_row(tbl2, list(row_data))

for i, width in enumerate([Cm(3.5), Cm(4.5), Cm(8.0)]):
    for row in tbl2.rows:
        row.cells[i].width = width

para('', space_before=6, space_after=0)

para('        Выбор SQLite в качестве СУБД обоснован следующим: приложение является '
     'настольным (desktop) и не требует сетевого сервера баз данных; SQLite не нуждается '
     'в отдельной установке; файл базы данных легко переносится и резервируется; '
     'производительность более чем достаточна для хранения диалогов и агентов.',
     indent=1.25)

para('        Выбор Newtonsoft.Json обусловлен широкой распространённостью библиотеки, '
     'гибкостью настройки сериализации (атрибуты, конвертеры) и совместимостью с '
     'форматами ответов всех используемых API-провайдеров.', indent=1.25)

# ─── 1.3 ──────────────────────────────────────────────────────
heading_left('1.3 Постановка задачи (техническое задание)')

para('        Требуется разработать настольное приложение FocusFlow AI — '
     'многопровайдерную AI-платформу на .NET 10 / Windows Forms. '
     'Приложение должно обеспечивать следующие возможности:')

tz_items = [
    'Подключение к пяти AI-провайдерам через REST API: Anthropic (Claude), OpenAI (GPT), '
     'Google (Gemini), Mistral, Groq (Llama). Ключи API хранятся в конфигурационном файле '
     'JSON в папке AppData пользователя.',
    'Три режима оркестрации: AUTO (интеллектуальный выбор модели), FUSION (параллельный '
     'вызов двух моделей + синтез), MANUAL (ручной выбор провайдера и модели).',
    'Система AI-агентов: 8 встроенных специализированных ассистентов + возможность '
     'создания и настройки пользовательских агентов.',
    'Хранение истории диалогов: таблицы Conversations и Messages в SQLite с поддержкой '
     'поиска, закрепления и удаления диалогов.',
    'Многошаговые воркфлоу: визуальный редактор последовательностей агентов с передачей '
     'контекста между шагами.',
    'Современный тёмный интерфейс с кастомными элементами управления (MessageBubble, '
     'AgentCard, RoundedPanel) и поддержкой Markdown-форматирования.',
    'Отмена запросов через CancellationToken.',
    'Автоматическое именование диалогов на основе первого сообщения.',
    'Конфигурирование: температура генерации, max_tokens, модель по умолчанию, '
     'режим оркестрации, количество сообщений истории в контексте.',
]
for item in tz_items:
    bullet(item)

para('        Нефункциональные требования: время отклика интерфейса не более 100 мс; '
     'все сетевые операции выполняются асинхронно без блокировки UI; '
     'хранение API-ключей только в локальном конфиге пользователя; '
     'поддержка отмены запроса в любой момент.')

# ─── 1.4 ──────────────────────────────────────────────────────
heading_left('1.4 Проектирование базы данных (ER-диаграмма, 3НФ)')

para('        Для хранения данных приложения используется SQLite. '
     'База данных содержит пять таблиц, нормализованных до третьей нормальной формы (3НФ).')

# Описание таблиц
tables_desc = [
    ('Conversations', 'Хранит метаданные диалогов (заголовок, агент, модель, дата создания, признак закрепления)'),
    ('Messages', 'Хранит отдельные сообщения диалога (роль, контент, дата, использованные токены, флаг ошибки)'),
    ('Agents', 'Описание AI-агентов (системный промпт, модель, настройки температуры, флаг встроенного агента)'),
    ('Workflows', 'Воркфлоу (название, описание, эмодзи, статус активности)'),
    ('WorkflowSteps', 'Шаги воркфлоу (порядок, агент, название шага, инструкция)'),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = 'Table Grid'
add_table_row(tbl3, ['Таблица', 'Назначение'], bold_row=True)
for row_data in tables_desc:
    add_table_row(tbl3, list(row_data))
for i, width in enumerate([Cm(4.5), Cm(11.5)]):
    for row in tbl3.rows:
        row.cells[i].width = width

para('', space_before=6, space_after=0)

heading_left('Структура таблиц')

# Conversations
para('Таблица Conversations:', bold=True, indent=0)
conv_cols = [
    ('Id', 'INTEGER', 'PK, AUTOINCREMENT', 'Первичный ключ'),
    ('Title', 'TEXT', 'NOT NULL', 'Заголовок диалога'),
    ('AgentId', 'TEXT', 'NOT NULL', 'Идентификатор агента (FK → Agents.Id)'),
    ('Model', 'TEXT', 'NOT NULL', 'Идентификатор модели'),
    ('IsPinned', 'INTEGER', 'NOT NULL DEFAULT 0', 'Флаг закрепления (0/1)'),
    ('CreatedAt', 'TEXT', 'NOT NULL', 'Дата создания (ISO 8601)'),
    ('UpdatedAt', 'TEXT', 'NOT NULL', 'Дата последнего обновления'),
]
tbl_c = doc.add_table(rows=1, cols=4)
tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_c.style = 'Table Grid'
add_table_row(tbl_c, ['Поле', 'Тип', 'Ограничение', 'Описание'], bold_row=True)
for row_data in conv_cols:
    add_table_row(tbl_c, list(row_data))
for i, w in enumerate([Cm(3.0), Cm(2.5), Cm(4.5), Cm(6.0)]):
    for row in tbl_c.rows:
        row.cells[i].width = w
para('', space_before=4, space_after=0)

# Messages
para('Таблица Messages:', bold=True, indent=0)
msg_cols = [
    ('Id', 'INTEGER', 'PK, AUTOINCREMENT', 'Первичный ключ'),
    ('ConversationId', 'INTEGER', 'FK → Conversations.Id ON DELETE CASCADE', 'Ссылка на диалог'),
    ('Role', 'TEXT', 'NOT NULL', 'Роль (user / assistant / system)'),
    ('Content', 'TEXT', 'NOT NULL', 'Текст сообщения'),
    ('CreatedAt', 'TEXT', 'NOT NULL', 'Дата создания'),
    ('TokensUsed', 'INTEGER', 'NOT NULL DEFAULT 0', 'Потреблённые токены'),
    ('ModelUsed', 'TEXT', 'NULL', 'Модель, использованная для ответа'),
    ('IsError', 'INTEGER', 'NOT NULL DEFAULT 0', 'Флаг сообщения об ошибке'),
]
tbl_m = doc.add_table(rows=1, cols=4)
tbl_m.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_m.style = 'Table Grid'
add_table_row(tbl_m, ['Поле', 'Тип', 'Ограничение', 'Описание'], bold_row=True)
for row_data in msg_cols:
    add_table_row(tbl_m, list(row_data))
for i, w in enumerate([Cm(3.5), Cm(2.0), Cm(5.0), Cm(5.5)]):
    for row in tbl_m.rows:
        row.cells[i].width = w
para('', space_before=4, space_after=0)

# Agents
para('Таблица Agents:', bold=True, indent=0)
ag_cols = [
    ('Id', 'TEXT', 'PK', 'Строковый идентификатор (default, code, writer …)'),
    ('Name', 'TEXT', 'NOT NULL', 'Отображаемое имя агента'),
    ('Description', 'TEXT', 'NOT NULL', 'Краткое описание'),
    ('SystemPrompt', 'TEXT', 'NOT NULL', 'Системный промпт'),
    ('Emoji', 'TEXT', 'NOT NULL', 'Иконка агента'),
    ('ColorHex', 'TEXT', 'NOT NULL', 'Цвет карточки (#RRGGBB)'),
    ('Model', 'TEXT', 'NOT NULL', 'Модель по умолчанию для агента'),
    ('Temperature', 'REAL', 'NOT NULL DEFAULT 0.7', 'Температура генерации'),
    ('MaxTokens', 'INTEGER', 'NOT NULL DEFAULT 4096', 'Максимальное число токенов'),
    ('IsBuiltIn', 'INTEGER', 'NOT NULL DEFAULT 0', 'Флаг встроенного агента'),
    ('CreatedAt', 'TEXT', 'NOT NULL', 'Дата создания'),
]
tbl_a = doc.add_table(rows=1, cols=4)
tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_a.style = 'Table Grid'
add_table_row(tbl_a, ['Поле', 'Тип', 'Ограничение', 'Описание'], bold_row=True)
for row_data in ag_cols:
    add_table_row(tbl_a, list(row_data))
for i, w in enumerate([Cm(3.0), Cm(2.0), Cm(4.5), Cm(6.5)]):
    for row in tbl_a.rows:
        row.cells[i].width = w
para('', space_before=4, space_after=0)

para('        Таблицы Workflows (Id, Name, Description, Emoji, IsActive, CreatedAt) и '
     'WorkflowSteps (Id, WorkflowId FK, StepOrder, AgentId, StepName, Instruction) '
     'организованы аналогично. WorkflowSteps связана с Workflows через FK ON DELETE CASCADE.')

heading_left('Соответствие нормальным формам (3НФ)')

nf_items = [
    '1НФ: все атрибуты атомарны (скалярные типы INTEGER, TEXT, REAL); составных атрибутов и групп нет.',
    '2НФ: все неключевые атрибуты полностью зависят от первичного ключа; составных ключей нет.',
    '3НФ: отсутствуют транзитивные зависимости — AgentId в Conversations — это внешний ключ, '
     'а не дублирование полей агента напрямую в таблице диалогов.',
]
for item in nf_items:
    bullet(item)

para('        Связи: Conversations (1) → Messages (N); Workflows (1) → WorkflowSteps (N). '
     'Связь Conversations → Agents — логическая (AgentId является строкой, '
     'равной Id в таблице Agents), что допустимо для SQLite.')

# ─── 1.5 ──────────────────────────────────────────────────────
heading_left('1.5 Проектирование интерфейса приложения (UI/UX)')

para('        Интерфейс приложения построен по принципу одностраничного приложения '
     '(SPA-like) внутри главного окна MainForm. Навигация осуществляется через '
     'боковую панель с кнопками-иконками. Все страницы реализованы как UserControl и '
     'переключаются программно без перезагрузки формы.')

heading_left('Формы и их назначение')

forms = [
    ('MainForm', 'Главная форма-контейнер. Содержит боковое меню навигации, статус-бар, '
     'область отображения текущей страницы. Хранит синглтон-ссылки на Config, Registry, Router, '
     'Orchestration, репозитории.'),
    ('ChatPage', 'Страница чата. Верхняя панель (заголовок, модель, агент-пикер, кнопки режимов AUTO/FUSION/MANUAL). '
     'Область сообщений (ScrollPanel с MessageBubble). Нижняя панель ввода (RichTextBox, кнопка Send/Stop, '
     'выбор провайдера/модели в ручном режиме). Анимированный индикатор "thinking".'),
    ('HistoryPage', 'Список всех диалогов с поиском, закреплением и удалением. '
     'Двойной клик открывает диалог в ChatPage.'),
    ('AgentsPage', 'Сетка карточек AgentCard. Кнопки «Новый агент», «Редактировать», «Удалить». '
     'Встроенные агенты не удаляются.'),
    ('AutomationPage', 'Редактор воркфлоу: список воркфлоу, редактор шагов, кнопка «Запустить». '
     'Прогресс выполнения отображается в текстовой области.'),
    ('SettingsPage', 'Настройки API-ключей (по одному полю на провайдера), выбор модели по умолчанию, '
     'режима оркестрации, температуры, max_tokens, числа сообщений в контексте.'),
]

tbl4 = doc.add_table(rows=1, cols=2)
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = 'Table Grid'
add_table_row(tbl4, ['Форма / Control', 'Назначение'], bold_row=True)
for row_data in forms:
    add_table_row(tbl4, list(row_data))
for i, w in enumerate([Cm(3.5), Cm(12.5)]):
    for row in tbl4.rows:
        row.cells[i].width = w

para('', space_before=6, space_after=0)

para('        Кастомные элементы управления: MessageBubble — пузырь сообщения с '
     'цветовым кодированием (фиолетовый — пользователь, тёмный — ассистент); '
     'AgentCard — карточка агента с эмодзи, цветной полосой и кнопками действий; '
     'RoundedPanel — панель со скруглёнными углами; '
     'InputDialog — диалог ввода строки с подтверждением.')

para('        Цветовая схема (тёмная): Background #1A1A2E, Surface #16213E, '
     'Card #0F3460, Accent #7C5CFC, Text #E0E0E0.')

# ─── 1.6 ──────────────────────────────────────────────────────
heading_left('1.6 Описание внутренней структуры приложения (архитектура)')

para('        Приложение построено по трёхслойной архитектуре с чётким разделением '
     'ответственности между слоями.')

layers = [
    ('Data (слой данных)', [
        'DatabaseManager — инициализация SQLite, создание таблиц, сидирование встроенных агентов.',
        'ConversationRepository — CRUD для таблицы Conversations (GetAll, GetById, Insert, Update, Delete, Search).',
        'MessageRepository — работа с Messages (GetByConversationId, Insert, DeleteByConversationId).',
        'AgentRepository — управление агентами (GetAll, GetById, Insert, Update, Delete).',
        'WorkflowRepository — воркфлоу и шаги (GetAll, Save, Delete, GetWithSteps).',
    ]),
    ('Services (слой бизнес-логики)', [
        'ILLMProvider — интерфейс провайдера: IsConfigured, SendAsync(systemPrompt, history, userMessage, model, temperature, maxTokens, ct).',
        'AnthropicProvider — реализация для Anthropic Claude API (формат messages с system-полем).',
        'OpenAICompatProvider — универсальная реализация для OpenAI-совместимых API (OpenAI, Mistral, Groq).',
        'GeminiProvider — реализация для Google Gemini API (специфический формат contents).',
        'ProviderRegistry — реестр провайдеров; GetAvailableModels() возвращает модели с настроенными ключами.',
        'AIRouter — классификация запроса (Code/Math/Creative/Analysis/Simple) и выбор оптимальной модели.',
        'OrchestrationService — выполнение запроса в режимах AUTO/FUSION/MANUAL.',
        'AIService — фасад для ChatPage; связывает репозитории и оркестратор.',
        'WorkflowService — последовательное выполнение шагов воркфлоу с передачей контекста.',
    ]),
    ('Forms + Controls (слой представления)', [
        'MainForm — главное окно, DI-root (создаёт и хранит все сервисы и репозитории).',
        'ChatPage / HistoryPage / AgentsPage / AutomationPage / SettingsPage — страницы навигации.',
        'MessageBubble / AgentCard / RoundedPanel / InputDialog — кастомные элементы управления.',
        'Theme — статический класс цветов и шрифтов.',
    ]),
]

for layer_name, items in layers:
    para(f'        {layer_name}:', bold=False, indent=1.25)
    for item in items:
        bullet(item, size=12)

para('        Взаимодействие слоёв: Forms → Services → Data. '
     'Forms не обращаются к Data напрямую. Services не знают о Forms. '
     'Data работает только с SQLite.')

page_break()

# ═════════════════════════════════════════════════════════════════
#  2. ПРАКТИЧЕСКАЯ ЧАСТЬ
# ═════════════════════════════════════════════════════════════════
heading('2. ПРАКТИЧЕСКАЯ ЧАСТЬ', size=14)

# ─── 2.1 ──────────────────────────────────────────────────────
heading_left('2.1 Создание проекта и подключение зависимостей')

para('        Проект создан в Visual Studio 2022 как «Windows Forms App (.NET)» '
     'с целевой платформой net10.0-windows. Файл проекта (.csproj) содержит '
     'следующие NuGet-зависимости:')

add_code_block(
    '<PackageReference Include="Newtonsoft.Json" Version="13.0.3" />\n'
    '<PackageReference Include="Microsoft.Data.Sqlite" Version="9.0.0" />'
)

para('        Пространства имён организованы по слоям:')
ns_items = [
    'FocusFlow_LMS.Models — модели данных (AppConfig, AIProvider, Conversation, Message, AIAgent, WorkflowStep);',
    'FocusFlow_LMS.Data — репозитории и DatabaseManager;',
    'FocusFlow_LMS.Services — провайдеры, роутер, оркестрация, воркфлоу;',
    'FocusFlow_LMS.Forms — страницы приложения;',
    'FocusFlow_LMS.Controls — кастомные элементы управления.',
]
for item in ns_items:
    bullet(item)

para('        Точка входа (Program.cs) настраивает Application.SetHighDpiMode, '
     'Application.EnableVisualStyles, устанавливает рендерер и запускает MainForm.')

# ─── 2.2 ──────────────────────────────────────────────────────
heading_left('2.2 Разработка базы данных (SQLite)')

para('        DatabaseManager.Initialize() вызывается при старте приложения и '
     'создаёт все таблицы с помощью команд CREATE TABLE IF NOT EXISTS. '
     'Путь к файлу базы данных:')

add_code_block(
    'Path.Combine(\n'
    '    Environment.GetFolderPath(Environment.SpecialFolder.ApplicationData),\n'
    '    "FocusFlowAI", "focusflow.db")'
)

para('        После создания таблиц вызывается SeedBuiltInAgents(), '
     'который добавляет 8 встроенных агентов командой INSERT OR IGNORE — '
     'то есть при повторном запуске агенты не дублируются.')

para('        Пример DDL-скрипта для таблицы Messages:')
add_code_block(
    'CREATE TABLE IF NOT EXISTS Messages (\n'
    '    Id             INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '    ConversationId INTEGER NOT NULL,\n'
    '    Role           TEXT    NOT NULL,\n'
    '    Content        TEXT    NOT NULL DEFAULT \'\',\n'
    '    CreatedAt      TEXT    NOT NULL,\n'
    '    TokensUsed     INTEGER NOT NULL DEFAULT 0,\n'
    '    ModelUsed      TEXT,\n'
    '    IsError        INTEGER NOT NULL DEFAULT 0,\n'
    '    FOREIGN KEY (ConversationId)\n'
    '        REFERENCES Conversations(Id) ON DELETE CASCADE\n'
    ');'
)

para('        Репозиторий ConversationRepository использует параметризованные запросы '
     '(AddWithValue) для защиты от SQL-инъекций. '
     'Все методы синхронны, так как SQLite-операции выполняются быстро '
     'и не требуют async-обёртки в данном контексте.')

# ─── 2.3 ──────────────────────────────────────────────────────
heading_left('2.3 Реализация основных функций (AI-провайдеры)')

para('        Интерфейс ILLMProvider определяет контракт для всех провайдеров:')

add_code_block(
    'public interface ILLMProvider\n'
    '{\n'
    '    bool IsConfigured { get; }\n'
    '    Task<AIResponse> SendAsync(\n'
    '        string            systemPrompt,\n'
    '        List<ChatMessage> history,\n'
    '        string            userMessage,\n'
    '        string            modelId,\n'
    '        float             temperature,\n'
    '        int               maxTokens,\n'
    '        CancellationToken ct);\n'
    '}'
)

para('        AnthropicProvider формирует тело запроса в формате Anthropic API '
     '(поле "system" выделено, "messages" содержит только user/assistant). '
     'OpenAICompatProvider использует единый формат OpenAI Chat Completions, '
     'что позволяет подключить OpenAI, Mistral и Groq через один класс с разными base URL. '
     'GeminiProvider формирует массив "contents" в формате Google Generative AI API.')

para('        Пример фрагмента OrchestrationService — режим FUSION:')

add_code_block(
    '// Параллельный вызов двух моделей\n'
    'var tasks = models.Take(2).Select(m =>\n'
    '    _registry.SendAsync(m, systemPrompt, history,\n'
    '        userMessage, temperature, maxTokens, ct)\n'
    ').ToList();\n'
    'var responses = await Task.WhenAll(tasks);\n\n'
    '// Синтез через самую дешёвую модель\n'
    'var synthModel = _registry.GetAvailableModels()\n'
    '    .Where(m => m.IsFast || m.CostPer1K < 0.002f)\n'
    '    .OrderBy(m => m.CostPer1K)\n'
    '    .FirstOrDefault() ?? models[0];'
)

para('        AIRouter классифицирует запрос методом Classify() на основе '
     'ключевых слов (на русском и английском языках). '
     'Функция Score() оценивает каждую модель по шкале QualityScore, '
     'бонусам за специализацию (IsCodeExpert, IsMath, IsCreative, IsAnalysis, IsFast) '
     'и штрафу за стоимость (CostPer1K).')

# ─── 2.4 ──────────────────────────────────────────────────────
heading_left('2.4 Разработка пользовательского интерфейса')

para('        ChatPage строится программно (без использования Windows Forms Designer) '
     'для полного контроля над макетом. Все элементы управления создаются '
     'и настраиваются в методах BuildTopBar(), BuildMessagesArea(), BuildInputArea(). '
     'Это обеспечивает независимость от .designer.cs-файлов и упрощает поддержку.')

para('        Отображение сообщений: при получении ответа от AI создаётся '
     'новый экземпляр MessageBubble и добавляется в ScrollPanel '
     'с прокруткой до нижнего края. Индикатор "обрабатывается..." '
     'реализован через System.Windows.Forms.Timer с анимацией точек (. → .. → ...).')

para('        Кнопки режимов AUTO / FUSION / MANUAL реализованы как ToggleButton '
     'с визуальной подсветкой активного режима (акцентный цвет #7C5CFC). '
     'Ручной режим (MANUAL) показывает дополнительный ряд с ComboBox выбора '
     'провайдера и модели.')

para('        AgentsPage отображает агентов в виде сетки карточек (FlowLayoutPanel). '
     'Каждая AgentCard содержит: цветную полосу (ColorHex), эмодзи, имя, описание, '
     'кнопки «Выбрать», «Редактировать», «Удалить» (для не-встроенных агентов).')

para('        HistoryPage содержит Panel с вертикальным списком диалогов '
     'и TextBox поиска. Поиск фильтрует диалоги по подстроке заголовка '
     'в реальном времени (событие TextChanged).')

# ─── 2.5 ──────────────────────────────────────────────────────
heading_left('2.5 Реализация системы оркестрации')

para('        OrchestrationService.RunAsync() принимает CancellationToken, '
     'что позволяет пользователю прервать запрос в любой момент нажатием кнопки Stop. '
     'Прогресс передаётся через IProgress<string>, что обеспечивает обновление '
     'UI-индикатора из фонового потока без явного Invoke.')

para('        WorkflowService.RunAsync() выполняет шаги последовательно в цикле. '
     'Результат (currentInput) каждого шага передаётся следующему. '
     'Если шаг задаёт явную инструкцию (step.Instruction), она предваряет '
     'входные данные; иначе входные данные передаются напрямую.')

para('        Конфигурация (AppConfig) сохраняется в JSON через Newtonsoft.Json:')

add_code_block(
    'public void Save()\n'
    '{\n'
    '    var dir = Path.GetDirectoryName(ConfigPath)!;\n'
    '    if (!Directory.Exists(dir)) Directory.CreateDirectory(dir);\n'
    '    File.WriteAllText(ConfigPath,\n'
    '        JsonConvert.SerializeObject(this, Formatting.Indented));\n'
    '}'
)

# ─── 2.6 ──────────────────────────────────────────────────────
heading_left('2.6 Обработка ошибок и валидация данных')

para('        Все сетевые вызовы обёрнуты в try/catch. '
     'В случае ошибки (HttpRequestException, TaskCanceledException, JsonException) '
     'провайдер возвращает AIResponse { Success = false, Error = "..." } '
     'вместо выброса исключения. Это предотвращает аварийное завершение приложения '
     'и позволяет отобразить понятное сообщение пользователю в виде MessageBubble '
     'с флагом IsError.')

para('        OrchestrationService реализует защитную логику:')
err_items = [
    'При FUSION, если оба провайдера вернули ошибку — возвращается сводное сообщение об ошибке;',
    'При FUSION, если один из двух провайдеров успешен — синтез пропускается, возвращается успешный ответ;',
    'При AUTO, если нет настроенных провайдеров — возвращается понятное сообщение «Добавьте API ключи»;',
    'При MANUAL, если модель не выбрана — возвращается «Выберите модель».',
]
for item in err_items:
    bullet(item)

para('        Валидация API-ключей: AppConfig.Load() и Save() используют try/catch '
     'с молчаливым возвратом нового экземпляра AppConfig() при любой ошибке чтения/записи. '
     'Это гарантирует запуск приложения даже при повреждённом конфиге.')

para('        Репозитории используют параметризованные SQL-запросы '
     '(cmd.Parameters.AddWithValue) для предотвращения SQL-инъекций. '
     'Все строки пользовательского ввода передаются как параметры, а не конкатенируются.')

page_break()

# ═════════════════════════════════════════════════════════════════
#  ЗАКЛЮЧЕНИЕ
# ═════════════════════════════════════════════════════════════════
heading('ЗАКЛЮЧЕНИЕ', size=14)

conclusion_texts = [
    ('        В ходе учебной практики по профессиональному модулю ПМ3 «Разработка '
     'модулей программного обеспечения для компьютерных систем» было разработано '
     'настольное приложение FocusFlow AI — многопровайдерная AI-платформа на базе '
     '.NET 10 и Windows Forms.'),
    ('        Реализовано полностью:'),
]
for t in conclusion_texts:
    para(t)

done_items = [
    'интеграция с пятью AI-провайдерами через REST API (Anthropic, OpenAI, Google, Mistral, Groq);',
    'система оркестрации в трёх режимах: AUTO, FUSION, MANUAL;',
    'интеллектуальный маршрутизатор AIRouter с классификацией запросов;',
    'полноценная CRUD-работа с диалогами, сообщениями, агентами и воркфлоу через SQLite;',
    'механизм многошаговой автоматизации (WorkflowService);',
    'современный тёмный интерфейс с кастомными элементами управления;',
    'асинхронная обработка запросов с поддержкой отмены (CancellationToken);',
    'конфигурирование приложения через JSON-файл в AppData.',
]
for item in done_items:
    bullet(item)

para('        Требует доработки / возможные улучшения:')

improve_items = [
    'добавление потокового вывода ответов (streaming) через Server-Sent Events;',
    'экспорт истории диалогов в форматы PDF и TXT;',
    'поддержка прикрепления файлов и изображений к сообщениям (vision API);',
    'локализация интерфейса (мультиязычность).',
]
for item in improve_items:
    bullet(item)

para('        В процессе выполнения практики были приобретены и закреплены следующие '
     'профессиональные компетенции: проектирование реляционных баз данных (3НФ) и '
     'работа с SQLite через Microsoft.Data.Sqlite; реализация асинхронного '
     'программирования (async/await, Task.WhenAll, CancellationToken, IProgress<T>); '
     'интеграция внешних REST API через HttpClient; архитектурное разделение '
     'приложения на слои (Data / Services / Forms); разработка событийно-ориентированного '
     'WinForms-интерфейса с кастомными элементами управления; обработка ошибок '
     'и безопасная работа с пользовательскими данными.')

para('        Поставленные цели практики достигнуты в полном объёме. '
     'Приложение FocusFlow AI представляет собой полностью работоспособный продукт, '
     'готовый к демонстрации и дальнейшему развитию.')

page_break()

# ═════════════════════════════════════════════════════════════════
#  СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
# ═════════════════════════════════════════════════════════════════
heading('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', size=14)

sources = [
    'Microsoft. C# documentation. – Режим доступа: https://learn.microsoft.com/en-us/dotnet/csharp/ (дата обращения: 04.2025).',
    'Microsoft. Windows Forms overview. – Режим доступа: https://learn.microsoft.com/en-us/dotnet/desktop/winforms/ (дата обращения: 04.2025).',
    'Microsoft. Asynchronous programming with async and await. – Режим доступа: https://learn.microsoft.com/en-us/dotnet/csharp/asynchronous-programming/ (дата обращения: 04.2025).',
    'Microsoft. Microsoft.Data.Sqlite overview. – Режим доступа: https://learn.microsoft.com/en-us/dotnet/standard/data/sqlite/ (дата обращения: 04.2025).',
    'Newtonsoft. Json.NET documentation. – Режим доступа: https://www.newtonsoft.com/json/help/html/Introduction.htm (дата обращения: 04.2025).',
    'Anthropic. API reference. Claude models. – Режим доступа: https://docs.anthropic.com/en/api/ (дата обращения: 04.2025).',
    'OpenAI. API reference. Chat completions. – Режим доступа: https://platform.openai.com/docs/api-reference/chat (дата обращения: 04.2025).',
    'Google. Gemini API documentation. – Режим доступа: https://ai.google.dev/api/ (дата обращения: 04.2025).',
    'Mistral AI. API documentation. – Режим доступа: https://docs.mistral.ai/api/ (дата обращения: 04.2025).',
    'Groq. API documentation. – Режим доступа: https://console.groq.com/docs/openai (дата обращения: 04.2025).',
    'SQLite. SQLite documentation. – Режим доступа: https://www.sqlite.org/docs.html (дата обращения: 04.2025).',
    'Microsoft. HttpClient class. – Режим доступа: https://learn.microsoft.com/en-us/dotnet/api/system.net.http.httpclient (дата обращения: 04.2025).',
    'Richter J. CLR via C#. 4th ed. – Redmond: Microsoft Press, 2012. – 896 с.',
    'Price M. C# 12 and .NET 8 – Modern Cross-Platform Development. – Packt Publishing, 2024. – 828 с.',
]

for i, source in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(-0.5)
    pf.left_indent = Cm(0.5)
    run = p.add_run(f'{i}. {source}')
    set_run_font(run, size=13)

page_break()

# ═════════════════════════════════════════════════════════════════
#  ПРИЛОЖЕНИЯ
# ═════════════════════════════════════════════════════════════════
heading('ПРИЛОЖЕНИЯ', size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf = p.paragraph_format
pf.space_before = Pt(0)
pf.space_after  = Pt(0)
r = p.add_run('Приложение А')
set_run_font(r, size=14, bold=True)

heading_left('Листинг 1 — AppConfig.cs (конфигурация приложения)')

add_code_block(
    'using Newtonsoft.Json;\n\n'
    'namespace FocusFlow_LMS.Models\n'
    '{\n'
    '    public class AppConfig\n'
    '    {\n'
    '        public string AnthropicApiKey { get; set; } = string.Empty;\n'
    '        public string OpenAiApiKey    { get; set; } = string.Empty;\n'
    '        public string GeminiApiKey    { get; set; } = string.Empty;\n'
    '        public string MistralApiKey   { get; set; } = string.Empty;\n'
    '        public string GroqApiKey      { get; set; } = string.Empty;\n\n'
    '        public string            DefaultModel         { get; set; } = "claude-opus-4-6";\n'
    '        public OrchestrationMode DefaultOrchestration { get; set; } = OrchestrationMode.Auto;\n'
    '        public int    MaxHistoryMessages { get; set; } = 10;\n'
    '        public float  Temperature        { get; set; } = 0.7f;\n'
    '        public int    MaxTokens          { get; set; } = 1536;\n'
    '        public bool   AutoTitleChats     { get; set; } = true;\n'
    '        public bool   ShowRouterInfo     { get; set; } = true;\n\n'
    '        private static readonly string ConfigPath = Path.Combine(\n'
    '            Environment.GetFolderPath(Environment.SpecialFolder.ApplicationData),\n'
    '            "FocusFlowAI", "config.json");\n\n'
    '        public static AppConfig Load()\n'
    '        {\n'
    '            try {\n'
    '                if (!File.Exists(ConfigPath)) return new AppConfig();\n'
    '                var json = File.ReadAllText(ConfigPath);\n'
    '                return JsonConvert.DeserializeObject<AppConfig>(json) ?? new AppConfig();\n'
    '            } catch { return new AppConfig(); }\n'
    '        }\n\n'
    '        public void Save()\n'
    '        {\n'
    '            try {\n'
    '                Directory.CreateDirectory(Path.GetDirectoryName(ConfigPath)!);\n'
    '                File.WriteAllText(ConfigPath,\n'
    '                    JsonConvert.SerializeObject(this, Formatting.Indented));\n'
    '            } catch { /* silent */ }\n'
    '        }\n'
    '    }\n'
    '}'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf = p.paragraph_format
pf.space_before = Pt(12)
pf.space_after  = Pt(0)
r = p.add_run('Приложение Б')
set_run_font(r, size=14, bold=True)

heading_left('Листинг 2 — OrchestrationService.cs (ключевые фрагменты)')

add_code_block(
    'public async Task<OrchestrationResult> RunAsync(\n'
    '    string systemPrompt, List<ChatMessage> history, string userMessage,\n'
    '    OrchestrationMode mode, ModelInfo? manualModel = null,\n'
    '    float temperature = 0.7f, int maxTokens = 4096,\n'
    '    IProgress<string>? progress = null, CancellationToken ct = default)\n'
    '{\n'
    '    var decision = _router.Route(userMessage, mode);\n\n'
    '    if (mode == OrchestrationMode.Fusion)\n'
    '    {\n'
    '        var tasks = decision.Models.Take(2)\n'
    '            .Select(m => _registry.SendAsync(m, systemPrompt, history,\n'
    '                userMessage, temperature, maxTokens, ct)).ToList();\n'
    '        var responses = await Task.WhenAll(tasks);\n'
    '        // ... synthesis ...\n'
    '    }\n'
    '}'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf = p.paragraph_format
pf.space_before = Pt(12)
pf.space_after  = Pt(0)
r = p.add_run('Приложение В')
set_run_font(r, size=14, bold=True)

heading_left('Листинг 3 — DatabaseManager.cs (фрагмент инициализации)')

add_code_block(
    'public static void Initialize()\n'
    '{\n'
    '    using var conn = new SqliteConnection(ConnectionString);\n'
    '    conn.Open();\n'
    '    var cmd = conn.CreateCommand();\n\n'
    '    cmd.CommandText = """\n'
    '        CREATE TABLE IF NOT EXISTS Conversations (\n'
    '            Id         INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '            Title      TEXT    NOT NULL DEFAULT \'Новый чат\',\n'
    '            AgentId    TEXT    NOT NULL DEFAULT \'default\',\n'
    '            Model      TEXT    NOT NULL DEFAULT \'claude-opus-4-6\',\n'
    '            IsPinned   INTEGER NOT NULL DEFAULT 0,\n'
    '            CreatedAt  TEXT    NOT NULL,\n'
    '            UpdatedAt  TEXT    NOT NULL\n'
    '        );\n'
    '        """;\n'
    '    cmd.ExecuteNonQuery();\n'
    '    // ... остальные таблицы ...\n'
    '    SeedBuiltInAgents(conn);\n'
    '}'
)

heading_left('Приложение Г — ER-диаграмма базы данных (текстовое представление)')

er_text = (
    'Conversations\n'
    '  Id (PK, INTEGER) ──────────────────────────────────────────────────────────────────────────┐\n'
    '  Title (TEXT)                                                                               │\n'
    '  AgentId (TEXT) ─→ Agents.Id (логическая связь)                                            │\n'
    '  Model (TEXT)                                                                               │\n'
    '  IsPinned (INTEGER)                                                     1:N (cascade delete)│\n'
    '  CreatedAt (TEXT)                                                                           │\n'
    '  UpdatedAt (TEXT)                                                                           │\n'
    '                                                                                             ▼\n'
    'Messages\n'
    '  Id (PK, INTEGER)\n'
    '  ConversationId (FK, INTEGER) ──────────────────────────────────────────────────────────────┘\n'
    '  Role (TEXT)\n'
    '  Content (TEXT)\n'
    '  CreatedAt (TEXT)\n'
    '  TokensUsed (INTEGER)\n'
    '  ModelUsed (TEXT)\n'
    '  IsError (INTEGER)\n\n'
    'Agents\n'
    '  Id (PK, TEXT)\n'
    '  Name, Description, SystemPrompt, Emoji, ColorHex, Model (TEXT)\n'
    '  Temperature (REAL)   MaxTokens (INTEGER)   IsBuiltIn (INTEGER)\n'
    '  CreatedAt (TEXT)\n\n'
    'Workflows\n'
    '  Id (PK, INTEGER) ─────────────────────────────────────────────────────────────────────────┐\n'
    '  Name, Description, Emoji (TEXT)                                        1:N (cascade delete)│\n'
    '  IsActive (INTEGER)   CreatedAt (TEXT)                                                      │\n'
    '                                                                                             ▼\n'
    'WorkflowSteps\n'
    '  Id (PK, INTEGER)\n'
    '  WorkflowId (FK, INTEGER) ─────────────────────────────────────────────────────────────────┘\n'
    '  StepOrder (INTEGER)   AgentId (TEXT) ─→ Agents.Id\n'
    '  StepName (TEXT)   Instruction (TEXT)'
)

p_er = doc.add_paragraph()
p_er.alignment = WD_ALIGN_PARAGRAPH.LEFT
pf_er = p_er.paragraph_format
pf_er.left_indent = Cm(0.5)
pf_er.space_before = Pt(4)
pf_er.space_after  = Pt(4)
pf_er.line_spacing_rule = WD_LINE_SPACING.SINGLE
run_er = p_er.add_run(er_text)
run_er.font.name = 'Courier New'
run_er.font.size = Pt(9)
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'),   'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'),  'F5F5F5')
p_er._p.get_or_add_pPr().append(shd2)

# ─────────────────────────────────────────────────────────────────
#  Сохранение
# ─────────────────────────────────────────────────────────────────
output_path = r"C:\Users\NOXQD\source\repos\FocusFlow LMS\Отчет_по_учебной_практике_FocusFlow_AI.docx"
doc.save(output_path)
print(f"Документ сохранён: {output_path}")
